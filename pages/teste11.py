import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os

# Conversion functions
def convert_pdf_to_word(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    
    doc = Document()
    doc.add_paragraph(text)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def convert_word_to_pdf(word_file):
    # Read the .docx file content
    doc = Document(word_file)
    text = "\n".join([p.text for p in doc.paragraphs])

    # Create a PDF
    output = BytesIO()
    c = canvas.Canvas(output, pagesize=letter)
    text_object = c.beginText(40, 750)  # Adjust for margins
    text_object.setFont("Helvetica", 12)

    # Split text into lines and add to PDF
    for line in text.splitlines():
        text_object.textLine(line)
    c.drawText(text_object)
    c.showPage()
    c.save()
    
    # Get PDF data
    output.seek(0)
    return output.getvalue()

# Streamlit app
st.sidebar.title('Navigation')
page = st.sidebar.radio('Go to', ['Introduction', 'Converter'])

# Path to image directory
image_dir = "./images"

# List of images with full paths
image_list = [os.path.join(image_dir, filename) for filename in [
    "File Converter.png",
    "File Converter 1.jpg",
    "File Converter 2.png",
    "File Converter 3.jpg"
]]

if page == 'Introduction':
    st.title('Welcome to the Universal File Converter App')
    
    with st.expander("Preview Images"):
        image_index = st.slider('Slider', 0, len(image_list) - 1)
        
        # Check if image path exists before displaying
        if os.path.exists(image_list[image_index]):
            st.image(image_list[image_index])
        else:
            st.error(f"Image not found: {image_list[image_index]}")
            st.write("Make sure the image files are located in the `images` folder.")
    
elif page == 'Converter':
    st.title('File Converter')
    
    conversion_type = st.selectbox("Select Conversion Type", ["PDF to Word", "Word to PDF"])
    
    if conversion_type == "PDF to Word":
        pdf_file = st.file_uploader("Upload PDF file", type=['pdf'])
        if pdf_file is not None:
            word_data = convert_pdf_to_word(pdf_file)
            st.download_button(label="Download Word file", data=word_data, file_name="converted.docx")
    
    elif conversion_type == "Word to PDF":
        word_file = st.file_uploader("Upload Word file", type=['docx'])
        if word_file is not None:
            pdf_data = convert_word_to_pdf(word_file)
            if pdf_data:
                st.download_button(label="Download PDF file", data=pdf_data, file_name="converted.pdf")
